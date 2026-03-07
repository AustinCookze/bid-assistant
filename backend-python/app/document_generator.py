"""
文档生成模块 - 生成投标文件 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Any, Optional
import io


class DocumentGenerator:
    """投标文件生成器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
    
    def generate(self, analysis_data: Dict[str, Any]) -> bytes:
        """生成投标文件"""
        # 1. 封面
        self._add_cover(analysis_data.get("project_info", {}))
        
        # 2. 评标索引
        if analysis_data.get("scoring_criteria"):
            self._add_scoring_index(analysis_data["scoring_criteria"])
        
        # 3. 资格证明
        self._add_qualification_section(analysis_data)
        
        # 4. 技术响应
        if analysis_data.get("technical_requirements"):
            self._add_technical_response(analysis_data["technical_requirements"])
        
        # 5. 商务响应
        self._add_business_response(analysis_data.get("business_terms", {}))
        
        # 6. 方案章节
        self._add_solution_sections(analysis_data.get("scoring_criteria", []))
        
        # 7. 政策章节
        self._add_policy_sections(analysis_data)
        
        # 8. 检查清单
        self._add_checklist(analysis_data.get("checklists", []))
        
        # 保存到字节流
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _add_cover(self, project_info: Dict[str, str]):
        """添加封面"""
        # 标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("投标文件")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        self.doc.add_paragraph()
        
        # 项目信息
        info_items = [
            ("项目名称", project_info.get("项目名称", "[USER_INPUT_项目名称]")),
            ("项目编号", project_info.get("项目编号", "[USER_INPUT_项目编号]")),
            ("投标人", "[USER_INPUT_投标人全称]"),
            ("日期", "[USER_INPUT_日期_YYYY年MM月DD日]"),
        ]
        
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}：{value}")
            run.font.size = Pt(16)
        
        self.doc.add_page_break()
    
    def _add_scoring_index(self, scoring_criteria: List[Dict]):
        """添加评标索引"""
        self._add_heading("评标索引表", level=1)
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        headers = ["序号", "评审因素", "分值", "投标文件页码"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 数据行
        for idx, item in enumerate(scoring_criteria, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.get("factor", "")[:50]
            row.cells[2].text = item.get("score", "")
            row.cells[3].text = "[PAGE_NUM_待填写]"
        
        self.doc.add_paragraph()
    
    def _add_qualification_section(self, analysis_data: Dict):
        """添加资格证明章节"""
        self._add_heading("一、资格证明文件", level=1)
        
        disqualification_items = analysis_data.get("disqualification_items", [])
        
        # 根据废标项动态生成
        categories = set(item.get("category", "") for item in disqualification_items)
        
        # 营业执照
        if any("营业执照" in str(item) for item in disqualification_items):
            self._add_heading("1.1 营业执照", level=2)
            self._add_placeholder("[请插入营业执照复印件，须加盖投标人公章]")
        
        # 财务报告
        if any(kw in str(disqualification_items) for kw in ["财务", "审计", "报表"]):
            self._add_heading("1.2 财务审计报告", level=2)
            self._add_placeholder("[请插入经审计的财务报告复印件]")
            self.doc.add_paragraph("注：如成立未满一年，可提供银行资信证明代替")
        
        # 纳税社保
        if any(kw in str(disqualification_items) for kw in ["纳税", "社保", "税收"]):
            self._add_heading("1.3 纳税和社保缴纳证明", level=2)
            self._add_placeholder("[请插入近期纳税凭据和社保缴纳证明]")
        
        # 信用查询
        if any(kw in str(disqualification_items) for kw in ["信用", "信用中国"]):
            self._add_heading("1.4 信用查询记录", level=2)
            self._add_placeholder("[请插入信用中国、中国政府采购网查询截图]")
        
        # 特定资质
        if any(kw in str(disqualification_items) for kw in ["ISO", "CMMI", "等保", "资质"]):
            self._add_heading("1.5 资质证书", level=2)
            self._add_placeholder("[请插入相关资质证书复印件]")
        
        # 人员资质
        if any(kw in str(disqualification_items) for kw in ["项目经理", "人员", "团队"]):
            self._add_heading("1.6 项目人员资质", level=2)
            self._add_placeholder("[请插入项目经理简历、证书及近X个月社保证明]")
    
    def _add_technical_response(self, technical_requirements: List[Dict]):
        """添加技术响应章节"""
        self._add_heading("二、技术响应文件", level=1)
        
        # 实质性条款响应表（★号参数）
        essential_params = [p for p in technical_requirements if p.get("is_essential")]
        if essential_params:
            self._add_heading("2.1 实质性技术条款响应表", level=2)
            self.doc.add_paragraph("⚠️ 注意：以下参数必须完全响应，不得负偏离，否则将导致废标！")
            self._add_tech_response_table(essential_params, is_essential=True)
        
        # 重要参数响应表（▲号参数）
        important_params = [p for p in technical_requirements if p.get("is_important")]
        if important_params:
            self._add_heading("2.2 重要技术参数响应表", level=2)
            self._add_tech_response_table(important_params)
        
        # 演示项清单（#号参数）
        demo_params = [p for p in technical_requirements if p.get("is_demo")]
        if demo_params:
            self._add_heading("2.3 演示项准备清单", level=2)
            self.doc.add_paragraph("⚠️ 本项目要求现场演示，请按以下清单准备：")
            for param in demo_params:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{param.get('param_name', '')}: {param.get('requirement', '')}")
            self.doc.add_paragraph()
            self.doc.add_paragraph("演示准备提示：").runs[0].font.bold = True
            self._add_placeholder("[USER_INPUT_演示环境准备说明]")
        
        # 一般参数响应表
        general_params = [p for p in technical_requirements 
                         if not p.get("is_essential") and not p.get("is_important") and not p.get("is_demo")]
        if general_params:
            self._add_heading("2.4 一般技术参数响应表", level=2)
            self._add_tech_response_table(general_params)
    
    def _add_tech_response_table(self, params: List[Dict], is_essential: bool = False):
        """添加技术响应表格"""
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        headers = ["序号", "参数名称", "招标要求", "投标响应", "证明材料"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for idx, param in enumerate(params, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = param.get("param_name", "")[:30]
            row.cells[2].text = param.get("requirement", "")[:100]
            
            if is_essential:
                row.cells[3].text = "[FIXED_完全响应，无偏离]"
            else:
                row.cells[3].text = "[USER_INPUT_响应说明]"
            
            row.cells[4].text = param.get("proof_required", "[证明材料]")
        
        self.doc.add_paragraph()
    
    def _add_business_response(self, business_terms: Dict):
        """添加商务响应章节"""
        self._add_heading("三、商务响应文件", level=1)
        
        # 交货期/工期承诺
        if business_terms.get("delivery_time"):
            self._add_heading("3.1 工期承诺", level=2)
            p = self.doc.add_paragraph()
            p.add_run("我公司承诺按照以下要求执行：\n")
            p.add_run(f"[FIXED_{business_terms.get('delivery_time', '按招标文件要求')}]")
        
        # 付款方式响应
        if business_terms.get("payment_terms"):
            self._add_heading("3.2 付款方式响应", level=2)
            self._add_placeholder("[响应招标文件要求的付款方式]")
        
        # 质保期承诺
        if business_terms.get("warranty_period"):
            self._add_heading("3.3 质保期承诺", level=2)
            p = self.doc.add_paragraph()
            p.add_run(f"质保期：[FIXED_{business_terms.get('warranty_period')}]")
        
        # 商务偏离表
        self._add_heading("3.4 商务条款偏离表", level=2)
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ["序号", "条款名称", "招标要求", "投标响应/偏离说明"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 至少添加一行无偏离
        row = table.add_row()
        row.cells[0].text = "1"
        row.cells[1].text = "综合"
        row.cells[2].text = "招标文件所有商务条款"
        row.cells[3].text = "[USER_INPUT_无偏离/偏离说明]"
        
        self.doc.add_paragraph()
    
    def _add_solution_sections(self, scoring_criteria: List[Dict]):
        """添加方案章节"""
        self._add_heading("四、技术方案", level=1)
        
        # 根据评分标准动态生成
        criteria_text = str(scoring_criteria)
        
        # 实施方案
        if any(kw in criteria_text for kw in ["实施", "项目方案", "进度"]):
            score = self._extract_score_for_factor(scoring_criteria, "实施")
            self._add_heading(f"4.1 项目实施方案 {score}", level=2)
            self._add_solution_template(["项目组织架构", "实施进度计划", "质量保障措施", "风险管理"])
        
        # 售后方案
        if any(kw in criteria_text for kw in ["售后", "服务", "培训"]):
            score = self._extract_score_for_factor(scoring_criteria, "服务")
            self._add_heading(f"4.2 售后服务方案 {score}", level=2)
            self._add_solution_template(["售后服务体系", "响应时间承诺", "技术支持方案", "培训计划"])
        
        # 设计方案
        if any(kw in criteria_text for kw in ["设计", "架构", "技术路线"]):
            score = self._extract_score_for_factor(scoring_criteria, "设计")
            self._add_heading(f"4.3 总体设计方案 {score}", level=2)
            self._add_solution_template(["系统架构设计", "技术路线选择", "关键技术说明", "安全保障设计"])
        
        # 业绩表
        if any(kw in criteria_text for kw in ["业绩", "案例", "经验"]):
            score = self._extract_score_for_factor(scoring_criteria, "业绩")
            self._add_heading(f"4.4 类似项目业绩 {score}", level=2)
            self.doc.add_paragraph("注：请注意招标文件对业绩时间范围的要求")
            self._add_performance_table()
    
    def _add_solution_template(self, sections: List[str]):
        """添加方案模板内容"""
        for section in sections:
            self._add_heading(section, level=3)
            self._add_placeholder(f"[USER_INPUT_{section}详细内容]")
    
    def _add_performance_table(self):
        """添加业绩表"""
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        headers = ["序号", "项目名称", "客户名称", "合同金额", "完成时间", "证明材料页码"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 添加示例行
        for i in range(1, 4):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = "[USER_INPUT_项目名称]"
            row.cells[2].text = "[USER_INPUT_客户名称]"
            row.cells[3].text = "[USER_INPUT_金额]"
            row.cells[4].text = "[USER_INPUT_时间_须在要求范围内]"
            row.cells[5].text = "[PAGE_NUM]"
        
        self.doc.add_paragraph()
    
    def _add_policy_sections(self, analysis_data: Dict):
        """添加政策章节"""
        raw_text = str(analysis_data)
        
        if any(kw in raw_text for kw in ["中小企业", "小微企业", "政策"]):
            self._add_heading("五、政府采购政策", level=1)
            
            if "中小企业" in raw_text or "小微企业" in raw_text:
                self._add_heading("5.1 中小企业声明函", level=2)
                self.doc.add_paragraph("根据《政府采购促进中小企业发展管理办法》，本公司声明如下：")
                self._add_placeholder("[请根据实际情况填写《中小企业声明函》]")
                self.doc.add_paragraph("注意：请准确填写所属行业类型，须与招标文件一致")
        else:
            self._add_heading("五、政府采购政策", level=1)
            self.doc.add_paragraph("本项目不适用政府采购优惠政策")
    
    def _add_checklist(self, checklists: List[str]):
        """添加检查清单"""
        self._add_heading("投标文件检查清单", level=1)
        
        self.doc.add_paragraph("在提交投标文件前，请逐项核对：").runs[0].font.bold = True
        
        for item in checklists:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        
        # 通用检查项
        self.doc.add_paragraph()
        self.doc.add_paragraph("通用检查项：").runs[0].font.bold = True
        general_checks = [
            "□ 投标文件是否按要求密封",
            "□ 正本/副本/电子版份数是否正确",
            "□ 签字盖章是否完整",
            "□ 投标保证金是否缴纳",
            "□ 投标有效期是否符合要求",
        ]
        for check in general_checks:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(check)
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return heading
    
    def _add_placeholder(self, text: str):
        """添加占位符段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
        run.font.italic = True
        return p
    
    def _extract_score_for_factor(self, scoring_criteria: List[Dict], keyword: str) -> str:
        """提取某项的分值"""
        for item in scoring_criteria:
            if keyword in item.get("factor", ""):
                score = item.get("score", "")
                return f"(本章节{score})" if score else ""
        return ""
