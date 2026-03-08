"""
文档解析模块 - 支持 PDF 和 Word 格式
使用 PyMuPDF 和 python-docx
"""
import fitz  # PyMuPDF
from docx import Document
import io
from typing import Optional
import re


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析 PDF 文件，返回纯文本"""
        text_parts = []
        
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                # 提取文本
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"\n--- 第{page_num}页 ---\n{text}")
                
                # 尝试提取表格内容（有些招标文件的表格是图片形式）
                tables = page.find_tables()
                if tables:
                    for table_idx, table in enumerate(tables.tables):
                        text_parts.append(f"\n[表格 {table_idx + 1}]")
                        for row in table.extract():
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    @staticmethod
    def parse_word(file_path: str) -> str:
        """解析 Word 文件（.docx），返回纯文本"""
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格内容
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n[表格 {table_idx + 1}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    @staticmethod
    def parse(file_path: str, file_type: Optional[str] = None) -> str:
        """根据文件类型自动选择解析方法"""
        if file_type is None:
            file_type = file_path.lower().split('.')[-1]
        
        if file_type in ['pdf']:
            return DocumentParser.parse_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return DocumentParser.parse_word(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_type}")


    @staticmethod
    def parse_from_bytes(content: bytes, file_type: str) -> str:
        """从字节流解析文档"""
        if file_type in ['pdf']:
            return DocumentParser._parse_pdf_bytes(content)
        elif file_type in ['docx', 'doc']:
            return DocumentParser._parse_word_bytes(content)
        else:
            raise ValueError(f"不支持的文件格式: {file_type}")
    
    @staticmethod
    def _parse_pdf_bytes(content: bytes) -> str:
        """从字节流解析 PDF"""
        text_parts = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"\n--- 第{page_num}页 ---\n{text}")
        return "\n".join(text_parts)
    
    @staticmethod
    def _parse_word_bytes(content: bytes) -> str:
        """从字节流解析 Word - 智能识别文档结构"""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # 检测并跳过目录页
        toc_patterns = [
            r'^\s*目\s*录\s*$',
            r'^\s*contents\s*$',
            r'^\s*第[一二三四五六七八九十\d]+章.*\.\.\.\.\.\.*\d+\s*$',  # 目录行：第一章......1
            r'^\s*\d+\.\d+.*\.\.\.\.\.\.*\d+\s*$',  # 目录行：1.1 标题......1
        ]
        
        # 页眉页脚常见模式
        header_footer_patterns = [
            r'^\s*第\s*\d+\s*页\s*共\s*\d+\s*页\s*$',
            r'^\s*Page\s*\d+\s*of\s*\d+\s*$',
            r'^\s*\d+\s*/\s*\d+\s*$',
        ]
        
        # 纯数字行（可能是页码）
        page_number_pattern = r'^\s*\d+\s*$'
        
        skip_count = 0
        max_skip = 50  # 最多跳过前50个段落（目录通常在前几页）
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 检测目录页并跳过
            if idx < max_skip:
                is_toc = False
                for pattern in toc_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        is_toc = True
                        skip_count += 1
                        break
                if is_toc:
                    continue
                
                # 如果连续跳过很多行，可能是目录，恢复正常处理
                if skip_count > 20 and idx > skip_count + 5:
                    pass  # 恢复正常处理
            
            # 跳过页眉页脚
            is_header_footer = False
            for pattern in header_footer_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    is_header_footer = True
                    break
            if is_header_footer:
                continue
            
            # 跳过孤立的页码数字
            if re.match(page_number_pattern, text) and len(text) < 5:
                continue
            
            # 检测真正的章节标题（用于结构化）
            # 区分：章节标题 vs 条款编号
            # 章节标题：第一章、第1章、一、（一）
            # 条款编号：5.8、5.8.1、（1）、1.
            
            # 真正的章节标题模式（大章节）
            chapter_patterns = [
                r'^第[一二三四五六七八九十百千]+章',  # 第一章、第十章
                r'^第\d+章',  # 第1章、第10章
                r'^[一二三四五六七八九十百千]+、',  # 一、二、
                r'^（[一二三四五六七八九十]）',  # （一）（二）
                r'^\([一二三四五六七八九十]\)',  # (一)(二)
            ]
            
            # 条款/条目编号（不是章节标题，保留在正文中）
            clause_patterns = [
                r'^\d+\.\d+\.\d+',  # 5.8.1, 1.1.1
                r'^\d+\.\d+',  # 5.8, 1.1
                r'^\d+\.',  # 1. 2.（但后面要跟内容）
                r'^（\d+）',  # （1）（2）
                r'^\(\d+\)',  # (1)(2)
            ]
            
            is_chapter = False
            for pattern in chapter_patterns:
                if re.match(pattern, text):
                    is_chapter = True
                    break
            
            # 如果是真正的章节标题，添加分隔
            if is_chapter:
                text_parts.append(f"\n{'='*40}\n{text}\n{'='*40}")
            else:
                # 条款编号保留原样，但添加换行便于阅读
                is_clause = False
                for pattern in clause_patterns:
                    if re.match(pattern, text):
                        is_clause = True
                        break
                
                if is_clause and not text_parts[-1].endswith('\n'):
                    text_parts.append('\n' + text)
                else:
                    text_parts.append(text)
        
        # 提取表格内容（评分标准、技术参数通常是表格）
        for table_idx, table in enumerate(doc.tables):
            # 检查表格内容，判断是否是内容表格还是布局表格
            table_texts = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    table_texts.append(row_text)
            
            # 如果表格有实质内容，保留
            if table_texts and len(table_texts) > 1:
                text_parts.append(f"\n[表格 {table_idx + 1}]")
                text_parts.extend(table_texts)
                text_parts.append("[表格结束]\n")
        
        return "\n".join(text_parts)


class TextCleaner:
    """文本清洗工具"""
    
    @staticmethod
    def clean(raw_text: str) -> str:
        """清洗提取的文本"""
        # 移除目录页常见的点线（如：第一章.......1）
        text = re.sub(r'\.{3,}\s*\d+\s*$', '', raw_text, flags=re.MULTILINE)
        
        # 移除多余空白行
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 移除页眉页脚常见的重复内容
        text = re.sub(r'第\s*\d+\s*页\s*共\s*\d+\s*页', '', text)
        
        # 移除孤立的数字（页码）
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        
        # 标准化空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 移除连续的重复行（可能是页眉）
        lines = text.split('\n')
        cleaned_lines = []
        prev_line = None
        for line in lines:
            line_stripped = line.strip()
            # 跳过与上一行完全相同的行（可能是重复的页眉）
            if line_stripped == prev_line and len(line_stripped) < 50:
                continue
            cleaned_lines.append(line)
            prev_line = line_stripped
        
        return '\n'.join(cleaned_lines).strip()
